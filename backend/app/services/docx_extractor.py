"""
DOCX text extractor — mirrors the pdf_extractor interface.
Uses python-docx to extract paragraph text from .docx files.
"""
import logging
from pathlib import Path
from typing import List, Tuple

logger = logging.getLogger("docx_extractor")


def extract_docx_detailed(path: Path) -> Tuple[str, int, List[Tuple[int, str]], bool, float]:
    """Extract text, estimated pages, page tuples, ocr_used, and ocr_confidence from a DOCX file."""
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx is not installed. DOCX text extraction skipped.")
        return "", 0, [], False, 1.0

    try:
        doc = Document(str(path))
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        paragraphs.append(cell_text)

        full_text = "\n".join(paragraphs)
        estimated_pages = max(1, len(full_text) // 3000) if full_text else 0

        # Build page tuples based on ~3000 char chunks or paragraphs
        page_tuples = []
        if full_text:
            chunk_size = 3000
            for page_idx in range(1, estimated_pages + 1):
                start = (page_idx - 1) * chunk_size
                end = start + chunk_size
                page_tuples.append((page_idx, full_text[start:end]))

        return full_text, estimated_pages, page_tuples, False, 1.0
    except Exception as e:
        logger.error("Failed to process DOCX %s: %s", path, e)
        return "", 0, [], False, 1.0


def extract_docx_info(path: Path) -> Tuple[str, int]:
    """Extract text and estimated page count from a DOCX file.
    Returns (text, estimated_page_count).
    """
    full_text, estimated_pages, _, _, _ = extract_docx_detailed(path)
    return full_text, estimated_pages


def extract_text_from_docx(path: Path) -> str:
    """Legacy wrapper — returns only text."""
    text, _ = extract_docx_info(path)
    return text

